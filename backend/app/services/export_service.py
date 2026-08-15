import io
import os
from typing import Literal

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# Polices avec support arabe, ordonnées par préférence
_ARABIC_FONT_CANDIDATES = [
    "Amiri-Regular.ttf",
    "DejaVuSans.ttf",
    "Arial Unicode MS.ttf",
    "NotoSansArabic-Regular.ttf",
    "Tahoma.ttf",
]


class ExportService:
    def __init__(self):
        self._arabic_font_name = None
        self._arabic_font_path = None

    def _find_arabic_font(self) -> tuple[str, str] | None:
        """Trouve une police TTF compatible arabe sur le système."""
        if self._arabic_font_name is not None:
            return self._arabic_font_name, self._arabic_font_path

        search_dirs = [
            "/usr/share/fonts",
            "/usr/local/share/fonts",
            "/System/Library/Fonts",
            "/Library/Fonts",
            "C:/Windows/Fonts",
            os.path.join(os.path.dirname(__file__), "..", "..", "assets", "fonts"),
        ]

        for directory in search_dirs:
            if not os.path.isdir(directory):
                continue
            for root, _, files in os.walk(directory):
                for candidate in _ARABIC_FONT_CANDIDATES:
                    if candidate in files:
                        path = os.path.join(root, candidate)
                        name = os.path.splitext(candidate)[0]
                        try:
                            pdfmetrics.registerFont(TTFont(name, path))
                            self._arabic_font_name = name
                            self._arabic_font_path = path
                            return name, path
                        except Exception:
                            continue
        return None

    @staticmethod
    def _is_arabic(text: str) -> bool:
        return any("\u0600" <= char <= "\u06FF" for char in text)

    @staticmethod
    def _reshape_arabic(text: str) -> str:
        """Reshape le texte arabe pour reportlab si les libs sont dispos."""
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
            reshaped = arabic_reshaper.reshape(text)
            return get_display(reshaped)
        except Exception:
            return text

    def export_text(self, text: str) -> bytes:
        return text.encode("utf-8")

    def export_docx(self, text: str, filename: str = "transcription") -> bytes:
        doc = DocxDocument()
        heading = doc.add_heading("Transcription Makhtout", level=1)

        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if self._is_arabic(text) else WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.runs[0]
        run.font.size = Pt(12)
        run.font.name = "Arial"

        # Métadonnées
        doc.add_paragraph()
        meta = doc.add_paragraph(f"Document : {filename}")
        meta.runs[0].font.size = Pt(9)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_pdf(self, text: str, filename: str = "transcription") -> bytes:
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        pdf.setTitle("Transcription Makhtout")

        font_info = self._find_arabic_font()
        body_font = font_info[0] if font_info else "Helvetica"
        title_font = body_font

        is_rtl = self._is_arabic(text)

        pdf.setFont(title_font, 16)
        pdf.drawString(50, height - 50, "Transcription Makhtout")

        pdf.setFont(body_font, 10)
        pdf.drawString(50, height - 70, f"Document : {filename}")

        # Texte avec retour à la ligne
        y = height - 110
        margin = 50
        max_width = width - 2 * margin
        line_height = 16
        font_size = 10

        pdf.setFont(body_font, font_size)

        for raw_line in text.split("\n"):
            line = self._reshape_arabic(raw_line) if is_rtl else raw_line
            words = line.split(" ")
            current_line = ""

            for word in words:
                test_line = current_line + " " + word if current_line else word
                if pdf.stringWidth(test_line, body_font, font_size) <= max_width:
                    current_line = test_line
                else:
                    if is_rtl:
                        pdf.drawRightString(width - margin, y, current_line)
                    else:
                        pdf.drawString(margin, y, current_line)
                    current_line = word
                    y -= line_height

                if y < margin:
                    pdf.showPage()
                    pdf.setFont(body_font, font_size)
                    y = height - margin

            if current_line:
                if is_rtl:
                    pdf.drawRightString(width - margin, y, current_line)
                else:
                    pdf.drawString(margin, y, current_line)
                y -= line_height

        pdf.save()
        buffer.seek(0)
        return buffer.getvalue()


export_service = ExportService()
